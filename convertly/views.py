from django.shortcuts import render, redirect
from django.contrib import messages
from django.http import HttpResponse
from .forms import DocumentForm  # Import the form
from pdf2docx import Converter
from docx2pdf import convert
from fpdf import FPDF
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image  
import xlwings as xw
import os




# Home view
def home(request):
    return render(request, 'home.html')


# Home view
def about(request):
    return render(request, 'about.html')



def search(request):
    messages.info(request, "I'm currently working on implementing a search feature to enhance your experience. Stay tuned for updates!")
    return render(request, "home.html")




# Success view to display the uploaded file
def success_page(request):
    # Retrieve file and format info from the session
    file = request.session.get('file')
    from_format = request.session.get('from_format')
    to_format = request.session.get('to_format')

    # Ensure these values exist
    if file:
        print(file)
        # want to get file name only before extension
        name = file.split('.')[0]
        print(name)

        # Define paths for the converted file
        original_file_path = os.path.join('uploads', file)  # Assuming you have an uploads directory
        converted_file_path = os.path.join('uploads', f'{name}.{to_format}')
        print(converted_file_path)

        # Conversion logic based on formats
        if from_format == 'pdf' and to_format == 'docx':
            convert_pdf_to_docx(original_file_path, converted_file_path)

        elif from_format == 'docx' and to_format == 'pdf':
            convert_docx_to_pdf(original_file_path, converted_file_path)

        elif from_format == 'txt' and to_format == 'pdf':
            convert_txt_to_pdf(original_file_path, converted_file_path)

        elif from_format == 'pdf' and to_format == 'txt':
            convert_pdf_to_txt(original_file_path, converted_file_path)

        

        elif from_format == 'png' and to_format == 'pdf':
            convert_png_to_pdf(original_file_path, converted_file_path)

            
        
        elif from_format == 'xlsx' and to_format == 'pdf':
            convert_xlsx_to_pdf(original_file_path, converted_file_path)
     
       

        
        elif from_format == 'docx' and to_format == 'txt':
            convert_docx_to_txt(original_file_path, converted_file_path)
        
        elif from_format == 'txt' and to_format == 'docx':
            convert_txt_to_docx(original_file_path, converted_file_path)
        
        
        elif from_format == 'jpg' and to_format == 'pdf':
            convert_jpg_to_pdf(original_file_path, converted_file_path)


        
        


        

        # Store the converted file name in the session
        request.session['converted_file_name'] = f'{name}.{to_format}'

        # Render the success template
        return render(request, 'success_page.html', {
            'file': file,
            'from_format': from_format,
            'to_format': to_format
        })
    else:
        # If no file in session, redirect back to home
        return redirect('home')
    





# Download file view
def download_file(request):
    # Retrieve the converted file name from the session
    file_name = request.session.get('converted_file_name')
    if not file_name:
        messages.error(request, 'No file to download.')
        return redirect('home')
    
    print(file_name)

    file_path = os.path.join('uploads', file_name)  # Adjust path based on your file storage
    
    print(file_path)

    # Serve the file for download
    with open(file_path, 'rb') as converted_file:
        response = HttpResponse(converted_file.read(), content_type='application/octet-stream')
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        return response







# Get document view
def get_document(request, from_format, to_format):
    form = DocumentForm()  # Initialize the form

    # Check if the request is a POST request
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)  # Bind data to the form
        if form.is_valid():
            # Get the file from the request
            file = request.FILES['file']

            # Check if the file is not empty
            if file:
                # Check if the file is of a valid format
                if file.name.endswith(from_format):
                    
                    if from_format == 'pdf' and to_format == 'jpg' or from_format=="pptx" and to_format=="pdf" or from_format=="html" and to_format=="pdf":
                        # wnat to return home page with error message
                        messages.info(request, f'I am currently working on the conversion from {from_format} to {to_format}. Thank you for your patience!"')
                        return redirect('home')
              


                    # Save file details in the session
                    request.session['file'] = file.name  # Save the file name
                    request.session['from_format'] = from_format
                    request.session['to_format'] = to_format

                    # Save the uploaded file to the uploads directory
                    with open(os.path.join('uploads', file.name), 'wb') as f:
                        for chunk in file.chunks():
                            f.write(chunk)

                    # Redirect to the success page
                    return redirect('success_page')
                else:
                    messages.error(request, f'Invalid file format. Please upload a file in {from_format} format')
            else:
                messages.error(request, 'File is empty')
        else:
            messages.error(request, 'Please correct the errors below')

    # Render form with context in the initial load or in case of errors
    return render(request, 'get_document.html', {
        'form': form,
        'from_format': from_format,
        'to_format': to_format
    })



# Conversion helper functions
def convert_pdf_to_docx(original_file_path, converted_file_path):
    """Convert PDF to DOCX."""
    cv = Converter(original_file_path)
    cv.convert(converted_file_path, start=0, end=None)
    cv.close()



import pythoncom


import os
import win32com.client

def convert_docx_to_pdf(input_path, output_path):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False

    # Ensure absolute paths
    abs_input = os.path.abspath(input_path)
    abs_output = os.path.abspath(output_path)

    # Open safely
    doc = word.Documents.Open(abs_input)
    doc.SaveAs(abs_output, FileFormat=17)  # 17 = PDF
    doc.Close()
    word.Quit()





def convert_txt_to_pdf(original_file_path, converted_file_path):
    """Convert TXT to PDF with proper text alignment."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    with open(original_file_path, 'r', encoding='utf-8') as txt_file:
        for line in txt_file:
            # Use multi_cell to handle long lines and ensure proper text alignment
            pdf.multi_cell(0, 10, txt=line.strip().encode('latin-1', 'replace').decode('latin-1'))

    pdf.output(converted_file_path)





def convert_pdf_to_txt(original_file_path, converted_file_path):
    """Convert PDF to TXT."""
    reader = PdfReader(original_file_path)
    with open(converted_file_path, 'w') as txt_file:
        for page in reader.pages:
            txt_file.write(page.extract_text())










def convert_docx_to_txt(original_file_path, converted_file_path):
    """Convert DOCX to TXT."""
    doc = Document(original_file_path)
    with open(converted_file_path, 'w', encoding='utf-8') as txt_file:
        for paragraph in doc.paragraphs:
            txt_file.write(paragraph.text + '\n')




def convert_txt_to_docx(original_file_path, converted_file_path):
    """Convert TXT to DOCX with improved formatting."""
    doc = Document()

    # Add a title for the document (you can customize this)
    doc.add_heading('Converted Document', level=1)

    # Open the TXT file and read its contents
    with open(original_file_path, 'r', encoding='utf-8') as txt_file:
        for line in txt_file:
          
            line = line.strip()
            if line:
               
                paragraph = doc.add_paragraph(line)
                
                if line.startswith('**'): 
                    paragraph.runs[0].bold = True
                    paragraph.text = paragraph.text[2:]  # Remove the '**' from the start
                elif line.startswith('*'): 
                    paragraph.runs[0].italic = True
                    paragraph.text = paragraph.text[1:]  # Remove the '*' from the start

                # Add a line break for better spacing (optional)
                doc.add_paragraph('')  # Adds an empty paragraph for spacing

    # Save the document
    doc.save(converted_file_path)




def convert_jpg_to_pdf(original_file_path, converted_file_path):
    """Convert JPG to PDF."""
    image = Image.open(original_file_path)
    image.convert('RGB').save(converted_file_path, "PDF")




def convert_xlsx_to_pdf(original_file_path, converted_file_path):
    """Convert XLSX to PDF using xlwings with original formatting."""
    
    # Open the Excel application
    app = xw.App(visible=False)
    
    try:
        # Open the workbook
        wb = app.books.open(original_file_path)

        # Save as PDF
        wb.to_pdf(converted_file_path)

        print(f"Excel file has been successfully converted to {converted_file_path}")
    
    finally:
        # Close the workbook and quit the app
        wb.close()
        app.quit()




# png to pdf
def convert_png_to_pdf(original_file_path, converted_file_path):
    """Convert PNG to PDF."""
    image = Image.open(original_file_path)
    image.convert('RGB').save(converted_file_path, "PDF")


